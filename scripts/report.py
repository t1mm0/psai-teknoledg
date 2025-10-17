#!/usr/bin/env python3
"""
PSAI_1 - Culture Current Lite: Report Generation Module
Purpose: Generate weekly auto-briefs with citations and analyst approval workflow
Last Modified: 2024-12-19 | By: AI Assistant | Completeness: 85/100
"""

import os
import sys
import json
import ollama
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from loguru import logger
import markdown
from docx import Document
from docx.shared import Inches
from jinja2 import Template

# Add project root to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

@dataclass
class ReportSection:
    """Report section structure"""
    title: str
    content: str
    insights: List[Dict]
    citations: List[str]
    confidence: float
    word_count: int

@dataclass
class WeeklyBrief:
    """Complete weekly brief structure"""
    title: str
    executive_summary: str
    sections: List[ReportSection]
    trends: List[Dict]
    key_findings: List[str]
    recommendations: List[str]
    generated_at: str
    metadata: Dict[str, Any]
    approval_status: str = "pending"
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None

class ReportGenerator:
    """Main report generation class"""
    
    def __init__(self, config_path: str = "config/report_config.json"):
        self.config = self._load_config(config_path)
        self.ollama_client = ollama.Client(host=self.config.get('ollama_host', 'http://localhost:11434'))
        
        logger.info("ReportGenerator initialized")
    
    def _load_config(self, config_path: str) -> Dict:
        """Load report configuration"""
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            logger.warning(f"Config file {config_path} not found, using defaults")
            return self._get_default_config()
    
    def _get_default_config(self) -> Dict:
        """Default configuration for report generation"""
        return {
            "report": {
                "title_template": "Culture Current Weekly Brief - {date}",
                "max_sections": 5,
                "max_words_per_section": 500,
                "include_executive_summary": True,
                "include_recommendations": True
            },
            "models": {
                "primary": "llama3.1",
                "fallback": "mistral",
                "temperature": 0.7,
                "max_tokens": 3000
            },
            "output": {
                "formats": ["markdown", "docx"],
                "save_path": "briefs/",
                "template_path": "templates/"
            },
            "approval": {
                "required": True,
                "approver_email": None,
                "approval_timeout_hours": 24
            }
        }
    
    def generate_weekly_brief(self, extraction_data: Dict) -> WeeklyBrief:
        """Generate complete weekly brief from extraction data"""
        logger.info("Starting weekly brief generation")
        
        insights = extraction_data.get('insights', [])
        trends = extraction_data.get('trends', [])
        
        # Generate report sections
        sections = self._generate_sections(insights, trends)
        
        # Generate executive summary
        executive_summary = self._generate_executive_summary(sections, trends)
        
        # Generate key findings
        key_findings = self._extract_key_findings(insights, trends)
        
        # Generate recommendations
        recommendations = self._generate_recommendations(sections, trends)
        
        # Create brief
        brief = WeeklyBrief(
            title=self._generate_title(),
            executive_summary=executive_summary,
            sections=sections,
            trends=trends,
            key_findings=key_findings,
            recommendations=recommendations,
            generated_at=datetime.now().isoformat(),
            metadata={
                'total_insights': len(insights),
                'total_trends': len(trends),
                'generation_model': self.config.get('models', {}).get('primary', 'llama3.1'),
                'config_version': '1.0'
            }
        )
        
        logger.info(f"Weekly brief generated: {len(sections)} sections, {len(key_findings)} findings")
        return brief
    
    def _generate_sections(self, insights: List[Dict], trends: List[Dict]) -> List[ReportSection]:
        """Generate report sections from insights and trends"""
        sections = []
        
        # Group insights by topic
        topic_groups = self._group_insights_by_topic(insights)
        
        # Create sections for each major topic
        for topic, topic_insights in topic_groups.items():
            if len(topic_insights) >= 2:  # Only create sections with sufficient content
                section = self._create_section(topic, topic_insights, trends)
                if section:
                    sections.append(section)
        
        # Limit number of sections
        max_sections = self.config.get('report', {}).get('max_sections', 5)
        sections = sections[:max_sections]
        
        return sections
    
    def _group_insights_by_topic(self, insights: List[Dict]) -> Dict[str, List[Dict]]:
        """Group insights by topic"""
        topics = {}
        
        # Topic keywords mapping
        topic_keywords = {
            'AI & Machine Learning': ['artificial intelligence', 'machine learning', 'AI', 'ML', 'neural', 'deep learning'],
            'Technology Trends': ['technology', 'tech', 'software', 'hardware', 'digital', 'innovation'],
            'Business & Markets': ['business', 'startup', 'company', 'market', 'investment', 'funding'],
            'Security & Privacy': ['security', 'privacy', 'cybersecurity', 'hack', 'breach', 'data protection'],
            'Emerging Technologies': ['blockchain', 'crypto', 'quantum', 'IoT', '5G', 'AR', 'VR']
        }
        
        for insight in insights:
            content_lower = (insight.get('title', '') + ' ' + insight.get('description', '')).lower()
            
            for topic, keywords in topic_keywords.items():
                if any(keyword in content_lower for keyword in keywords):
                    if topic not in topics:
                        topics[topic] = []
                    topics[topic].append(insight)
                    break
        
        return topics
    
    def _create_section(self, topic: str, insights: List[Dict], trends: List[Dict]) -> Optional[ReportSection]:
        """Create a report section for a topic"""
        # Find relevant trends
        relevant_trends = [t for t in trends if topic.lower() in t.get('trend_name', '').lower()]
        
        # Generate section content using Ollama
        content = self._generate_section_content(topic, insights, relevant_trends)
        
        if not content:
            return None
        
        # Extract citations
        citations = self._extract_citations_from_insights(insights)
        
        # Calculate confidence
        confidence = sum(insight.get('confidence', 0.5) for insight in insights) / len(insights)
        
        return ReportSection(
            title=topic,
            content=content,
            insights=insights,
            citations=citations,
            confidence=confidence,
            word_count=len(content.split())
        )
    
    def _generate_section_content(self, topic: str, insights: List[Dict], trends: List[Dict]) -> str:
        """Generate section content using Ollama"""
        prompt = f"""
You are a professional business analyst writing a weekly intelligence brief.

Topic: {topic}

Key Insights:
{self._format_insights_for_prompt(insights)}

Relevant Trends:
{self._format_trends_for_prompt(trends)}

Write a comprehensive section (400-500 words) that:
1. Summarizes the key developments in this topic area
2. Explains the significance and implications
3. Highlights any notable trends or patterns
4. Uses a professional, analytical tone
5. Includes specific examples and data points

Focus on actionable insights and clear explanations.
"""
        
        try:
            response = self._call_ollama(prompt)
            return response.strip()
        except Exception as e:
            logger.error(f"Section content generation failed for {topic}: {e}")
            return f"Analysis of {topic} based on {len(insights)} insights and {len(trends)} trends."
    
    def _generate_executive_summary(self, sections: List[ReportSection], trends: List[Dict]) -> str:
        """Generate executive summary"""
        prompt = f"""
You are writing an executive summary for a weekly intelligence brief.

Report Sections:
{self._format_sections_for_prompt(sections)}

Key Trends:
{self._format_trends_for_prompt(trends)}

Write a concise executive summary (200-300 words) that:
1. Highlights the most important developments
2. Identifies key themes and patterns
3. Provides strategic implications
4. Uses executive-level language
5. Focuses on high-impact insights

Be direct, actionable, and strategic in your analysis.
"""
        
        try:
            response = self._call_ollama(prompt)
            return response.strip()
        except Exception as e:
            logger.error(f"Executive summary generation failed: {e}")
            return f"Executive summary covering {len(sections)} key areas and {len(trends)} major trends."
    
    def _extract_key_findings(self, insights: List[Dict], trends: List[Dict]) -> List[str]:
        """Extract key findings from insights and trends"""
        findings = []
        
        # Extract high-confidence insights
        high_confidence_insights = [i for i in insights if i.get('confidence', 0) > 0.7]
        
        for insight in high_confidence_insights[:5]:  # Top 5 insights
            findings.append(f"{insight.get('title', 'Key Insight')}: {insight.get('description', '')[:100]}...")
        
        # Extract trend findings
        for trend in trends[:3]:  # Top 3 trends
            findings.append(f"Trend: {trend.get('trend_name', 'Unknown')} is {trend.get('momentum', 'stable')} with {trend.get('confidence', 0.5):.1%} confidence")
        
        return findings
    
    def _generate_recommendations(self, sections: List[ReportSection], trends: List[Dict]) -> List[str]:
        """Generate actionable recommendations"""
        prompt = f"""
Based on the following intelligence analysis, provide 3-5 actionable recommendations for business leaders.

Report Sections:
{self._format_sections_for_prompt(sections)}

Key Trends:
{self._format_trends_for_prompt(trends)}

Provide specific, actionable recommendations that:
1. Address the key trends and developments
2. Are practical and implementable
3. Consider both opportunities and risks
4. Are relevant for business strategy
5. Include timeframes where appropriate

Format as a numbered list of recommendations.
"""
        
        try:
            response = self._call_ollama(prompt)
            # Parse numbered list
            recommendations = []
            for line in response.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-')):
                    recommendations.append(line)
            return recommendations[:5]  # Limit to 5 recommendations
        except Exception as e:
            logger.error(f"Recommendations generation failed: {e}")
            return ["Monitor key trends identified in this brief", "Review strategic implications of emerging technologies"]
    
    def _generate_title(self) -> str:
        """Generate report title"""
        date_str = datetime.now().strftime("%B %d, %Y")
        template = self.config.get('report', {}).get('title_template', "Culture Current Weekly Brief - {date}")
        return template.format(date=date_str)
    
    def _extract_citations_from_insights(self, insights: List[Dict]) -> List[str]:
        """Extract citations from insights"""
        citations = []
        
        for insight in insights:
            if insight.get('insight_type') == 'citation':
                citations.append(insight.get('source_url', ''))
            elif insight.get('source_url'):
                citations.append(insight.get('source_url'))
        
        # Remove duplicates and limit
        citations = list(set(citations))[:10]
        return citations
    
    def _format_insights_for_prompt(self, insights: List[Dict]) -> str:
        """Format insights for prompt"""
        formatted = []
        for insight in insights[:5]:  # Limit to 5 insights
            formatted.append(f"- {insight.get('title', 'Insight')}: {insight.get('description', '')[:200]}...")
        return '\n'.join(formatted)
    
    def _format_trends_for_prompt(self, trends: List[Dict]) -> str:
        """Format trends for prompt"""
        formatted = []
        for trend in trends:
            formatted.append(f"- {trend.get('trend_name', 'Trend')}: {trend.get('description', '')} (Momentum: {trend.get('momentum', 'stable')})")
        return '\n'.join(formatted)
    
    def _format_sections_for_prompt(self, sections: List[ReportSection]) -> str:
        """Format sections for prompt"""
        formatted = []
        for section in sections:
            formatted.append(f"- {section.title}: {section.content[:200]}...")
        return '\n'.join(formatted)
    
    def _call_ollama(self, prompt: str) -> str:
        """Call Ollama API with error handling"""
        model = self.config.get('models', {}).get('primary', 'llama3.1')
        fallback_model = self.config.get('models', {}).get('fallback', 'mistral')
        
        try:
            response = self.ollama_client.generate(
                model=model,
                prompt=prompt,
                options={
                    'temperature': self.config.get('models', {}).get('temperature', 0.7),
                    'num_predict': self.config.get('models', {}).get('max_tokens', 3000)
                }
            )
            return response['response']
            
        except Exception as e:
            logger.warning(f"Primary model {model} failed, trying fallback: {e}")
            
            try:
                response = self.ollama_client.generate(
                    model=fallback_model,
                    prompt=prompt,
                    options={
                        'temperature': self.config.get('models', {}).get('temperature', 0.7),
                        'num_predict': self.config.get('models', {}).get('max_tokens', 3000)
                    }
                )
                return response['response']
                
            except Exception as e2:
                logger.error(f"Both models failed: {e2}")
                raise
    
    def save_markdown(self, brief: WeeklyBrief, filename: str = None) -> str:
        """Save brief as Markdown"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"briefs/weekly_brief_{timestamp}.md"
        
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        # Load template
        template_path = os.path.join(self.config.get('output', {}).get('template_path', 'templates/'), 'brief_template.md')
        try:
            with open(template_path, 'r') as f:
                template = Template(f.read())
        except FileNotFoundError:
            template = Template(self._get_default_markdown_template())
        
        # Render template
        content = template.render(
            title=brief.title,
            executive_summary=brief.executive_summary,
            sections=brief.sections,
            trends=brief.trends,
            key_findings=brief.key_findings,
            recommendations=brief.recommendations,
            generated_at=brief.generated_at,
            metadata=brief.metadata
        )
        
        # Save file
        with open(filename, 'w') as f:
            f.write(content)
        
        logger.info(f"Markdown brief saved to {filename}")
        return filename
    
    def save_docx(self, brief: WeeklyBrief, filename: str = None) -> str:
        """Save brief as DOCX"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"briefs/weekly_brief_{timestamp}.docx"
        
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(brief.title, 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(brief.executive_summary)
        
        # Sections
        for section in brief.sections:
            doc.add_heading(section.title, level=1)
            doc.add_paragraph(section.content)
            
            # Add citations if any
            if section.citations:
                doc.add_heading('Sources', level=2)
                for citation in section.citations:
                    doc.add_paragraph(citation, style='List Bullet')
        
        # Key Findings
        if brief.key_findings:
            doc.add_heading('Key Findings', level=1)
            for finding in brief.key_findings:
                doc.add_paragraph(finding, style='List Bullet')
        
        # Recommendations
        if brief.recommendations:
            doc.add_heading('Recommendations', level=1)
            for recommendation in brief.recommendations:
                doc.add_paragraph(recommendation, style='List Number')
        
        # Metadata
        doc.add_heading('Report Metadata', level=1)
        doc.add_paragraph(f"Generated: {brief.generated_at}")
        doc.add_paragraph(f"Total Insights: {brief.metadata.get('total_insights', 0)}")
        doc.add_paragraph(f"Total Trends: {brief.metadata.get('total_trends', 0)}")
        doc.add_paragraph(f"Model: {brief.metadata.get('generation_model', 'Unknown')}")
        
        # Save document
        doc.save(filename)
        
        logger.info(f"DOCX brief saved to {filename}")
        return filename
    
    def _get_default_markdown_template(self) -> str:
        """Default Markdown template"""
        return """# {{ title }}

**Generated:** {{ generated_at }}

## Executive Summary

{{ executive_summary }}

## Key Findings

{% for finding in key_findings %}
- {{ finding }}
{% endfor %}

## Report Sections

{% for section in sections %}
### {{ section.title }}

{{ section.content }}

**Sources:**
{% for citation in section.citations %}
- {{ citation }}
{% endfor %}

{% endfor %}

## Trends Analysis

{% for trend in trends %}
### {{ trend.trend_name }}
- **Momentum:** {{ trend.momentum }}
- **Confidence:** {{ trend.confidence }}
- **Description:** {{ trend.description }}

{% endfor %}

## Recommendations

{% for recommendation in recommendations %}
{{ loop.index }}. {{ recommendation }}
{% endfor %}

## Report Metadata

- **Total Insights:** {{ metadata.total_insights }}
- **Total Trends:** {{ metadata.total_trends }}
- **Generation Model:** {{ metadata.generation_model }}
- **Config Version:** {{ metadata.config_version }}
"""

def main():
    """Main execution function"""
    import argparse
    
    parser = argparse.ArgumentParser(description='PSAI_1 Report Generator')
    parser.add_argument('--input', required=True, help='Input extraction data file')
    parser.add_argument('--output', help='Output directory')
    parser.add_argument('--format', choices=['markdown', 'docx', 'both'], default='both', help='Output format')
    parser.add_argument('--test', action='store_true', help='Run in test mode')
    
    args = parser.parse_args()
    
    # Setup logging
    logger.add("logs/report.log", rotation="1 week", retention="1 month")
    
    # Load extraction data
    with open(args.input, 'r') as f:
        extraction_data = json.load(f)
    
    if args.test:
        logger.info("Running in test mode")
        # Limit data for testing
        extraction_data['insights'] = extraction_data.get('insights', [])[:10]
        extraction_data['trends'] = extraction_data.get('trends', [])[:3]
    
    # Initialize generator
    generator = ReportGenerator()
    
    # Generate brief
    brief = generator.generate_weekly_brief(extraction_data)
    
    # Save in requested formats
    output_dir = args.output or "briefs/"
    
    if args.format in ['markdown', 'both']:
        md_file = generator.save_markdown(brief, os.path.join(output_dir, "weekly_brief.md"))
        print(f"Markdown brief saved: {md_file}")
    
    if args.format in ['docx', 'both']:
        docx_file = generator.save_docx(brief, os.path.join(output_dir, "weekly_brief.docx"))
        print(f"DOCX brief saved: {docx_file}")
    
    print(f"Report generation complete: {len(brief.sections)} sections, {len(brief.key_findings)} findings")
    logger.info(f"Report generation complete: {len(brief.sections)} sections, {len(brief.key_findings)} findings")

if __name__ == "__main__":
    main()
